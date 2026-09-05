from __future__ import annotations

import argparse
import json
from datetime import datetime
from pathlib import Path
from typing import Any, Iterable

from finance_agent.excel_voucher_parser import (
    archive_source_file,
    derive_direction_and_amount,
    is_blank_row,
    make_record_id,
    normalize_date,
    normalize_text,
    parse_attachment_note,
    parse_budget_category,
    parse_decimal,
    parse_project_code,
    parse_reference_numbers,
)
from finance_agent.excel_voucher_parser import parse_workbook as parse_excel_workbook
from finance_agent.voucher_record import SourceCell, VoucherRecord, decimal_to_money


SUPPORTED_SUFFIXES = {".xlsx", ".xlsm", ".docx", ".pdf"}

FIELD_ALIASES: dict[str, tuple[str, ...]] = {
    "voucher_date": ("日期", "凭证日期", "付款日期", "交易日期", "回单日期"),
    "voucher_no": ("凭证编号", "凭证号", "凭证字号", "记账凭证号"),
    "summary": ("摘要", "用途", "事由", "付款摘要", "说明"),
    "income_amount": ("收入/拨入", "收入", "拨入", "借方金额"),
    "expense_amount": ("支出/拨出", "支出", "拨出", "贷方金额", "付款金额", "金额"),
    "budget_category": ("预算分类", "预算科目", "经费类别", "支出类别", "费用类别"),
    "attachment_note": ("附件材料", "附件", "备注", "项目经费号", "经费号", "负责人"),
}


def parse_financial_document(
    input_path: Path,
    raw_archive_dir: Path | None = None,
    include_debug: bool = False,
) -> dict[str, Any]:
    suffix = input_path.suffix.lower()
    if suffix in {".xlsx", ".xlsm"}:
        return parse_excel_workbook(input_path, raw_archive_dir=raw_archive_dir, include_debug=include_debug)
    if suffix == ".docx":
        return parse_docx(input_path, raw_archive_dir=raw_archive_dir, include_debug=include_debug)
    if suffix == ".pdf":
        return parse_pdf(input_path, raw_archive_dir=raw_archive_dir, include_debug=include_debug)
    raise ValueError(f"Unsupported financial document type: {input_path.suffix}. Supported: {sorted(SUPPORTED_SUFFIXES)}")


def parse_docx(
    input_path: Path,
    raw_archive_dir: Path | None = None,
    include_debug: bool = False,
) -> dict[str, Any]:
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError("Missing dependency: install python-docx with `python -m pip install -r requirements.txt`.") from exc

    input_path = input_path.resolve()
    archived_source = archive_source_file(input_path, raw_archive_dir) if raw_archive_dir else None
    document = Document(str(input_path))

    records: list[VoucherRecord] = []
    table_reports: list[dict[str, Any]] = []
    for table_index, table in enumerate(document.tables, start=1):
        rows = [[normalize_text(cell.text) for cell in row.cells] for row in table.rows]
        report, table_records = parse_table_rows(rows, input_path, f"docx_table_{table_index}")
        table_reports.append(report)
        records.extend(table_records)

    payload = build_payload(input_path, archived_source, records, table_reports, include_debug)
    if include_debug:
        payload["debug"]["paragraph_count"] = len(document.paragraphs)
    return payload


def parse_pdf(
    input_path: Path,
    raw_archive_dir: Path | None = None,
    include_debug: bool = False,
) -> dict[str, Any]:
    input_path = input_path.resolve()
    archived_source = archive_source_file(input_path, raw_archive_dir) if raw_archive_dir else None
    table_reports: list[dict[str, Any]] = []
    records: list[VoucherRecord] = []

    try:
        import pdfplumber

        with pdfplumber.open(str(input_path)) as pdf:
            for page_index, page in enumerate(pdf.pages, start=1):
                page_extracted = False
                for table_index, table in enumerate(page.extract_tables() or [], start=1):
                    rows = [[normalize_text(cell) for cell in row] for row in table if row]
                    report, table_records = parse_table_rows(rows, input_path, f"pdf_page_{page_index}_table_{table_index}")
                    table_reports.append(report)
                    records.extend(table_records)
                    page_extracted = True
                if not page_extracted:
                    text_rows = text_to_rows(page.extract_text() or "")
                    report, table_records = parse_table_rows(text_rows, input_path, f"pdf_page_{page_index}_text")
                    table_reports.append(report)
                    records.extend(table_records)
    except ImportError:
        records, table_reports = parse_pdf_with_pypdf(input_path)

    return build_payload(input_path, archived_source, records, table_reports, include_debug)


def parse_pdf_with_pypdf(input_path: Path) -> tuple[list[VoucherRecord], list[dict[str, Any]]]:
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise RuntimeError("Missing dependency: install pdfplumber or pypdf with `python -m pip install -r requirements.txt`.") from exc

    reader = PdfReader(str(input_path))
    records: list[VoucherRecord] = []
    reports: list[dict[str, Any]] = []
    for page_index, page in enumerate(reader.pages, start=1):
        text_rows = text_to_rows(page.extract_text() or "")
        report, page_records = parse_table_rows(text_rows, input_path, f"pdf_page_{page_index}_text")
        reports.append(report)
        records.extend(page_records)
    return records, reports


def text_to_rows(text: str) -> list[list[str | None]]:
    rows: list[list[str | None]] = []
    for line in text.splitlines():
        value = normalize_text(line)
        if not value:
            continue
        if "\t" in value:
            rows.append([normalize_text(part) for part in value.split("\t")])
        else:
            import re

            rows.append([normalize_text(part) for part in re.split(r"\s{2,}", value) if normalize_text(part)])
    return rows


def parse_table_rows(
    rows: list[list[Any]],
    source_file: Path,
    source_name: str,
) -> tuple[dict[str, Any], list[VoucherRecord]]:
    if not rows:
        return empty_report(source_name), []

    header_row_index = detect_header_row(rows)
    headers = [normalize_text(value) for value in rows[header_row_index]]
    field_map = map_headers(headers)
    records: list[VoucherRecord] = []

    if not field_map:
        return {
            **empty_report(source_name),
            "row_count": len(rows),
            "header_row": None,
            "headers": headers,
            "field_map": {},
            "warning": "no supported voucher columns detected",
        }, []

    for row_index, row in enumerate(rows[header_row_index + 1 :], start=header_row_index + 2):
        if is_blank_row(row):
            continue
        normalized_row = [normalize_text(value) for value in row]
        raw_by_header = build_raw_row(headers, normalized_row)
        warnings: list[str] = []

        voucher_date = normalize_date(get_field_value(field_map, "voucher_date", normalized_row), warnings)
        voucher_no = normalize_text(get_field_value(field_map, "voucher_no", normalized_row))
        summary = normalize_text(get_field_value(field_map, "summary", normalized_row))
        income = parse_decimal(get_field_value(field_map, "income_amount", normalized_row), "income_amount", warnings)
        expense = parse_decimal(get_field_value(field_map, "expense_amount", normalized_row), "expense_amount", warnings)
        direction, signed_amount = derive_direction_and_amount(income, expense, warnings)
        budget_code, budget_name = parse_budget_category(get_field_value(field_map, "budget_category", normalized_row))
        attachment_note = normalize_text(get_field_value(field_map, "attachment_note", normalized_row))
        fund_owner, project_fund_no = parse_attachment_note(attachment_note)
        project_code = parse_project_code(summary)
        references = parse_reference_numbers(summary)

        if not expense and not income:
            warnings.append("no amount column value found")

        records.append(
            VoucherRecord(
                record_id=make_record_id(source_file, source_name, row_index, voucher_no, voucher_date),
                source_file=str(source_file),
                source_sheet=source_name,
                source_row=row_index,
                voucher_date=voucher_date,
                voucher_no=voucher_no,
                summary=summary,
                income_amount=decimal_to_money(income),
                expense_amount=decimal_to_money(expense),
                signed_amount=decimal_to_money(signed_amount),
                direction=direction,
                budget_category_code=budget_code,
                budget_category_name=budget_name,
                project_code=project_code,
                project_fund_no=project_fund_no,
                fund_owner=fund_owner,
                bank_receipt_reference_numbers=references,
                attachment_note=attachment_note,
                source_cells=build_source_cells(source_name, headers, normalized_row, field_map, row_index),
                raw_row=raw_by_header,
                warnings=warnings,
            )
        )

    return {
        "source": source_name,
        "row_count": len(rows),
        "header_row": header_row_index + 1,
        "headers": headers,
        "field_map": {field: {"column": idx + 1, "header": headers[idx] if idx < len(headers) else None} for field, idx in field_map.items()},
        "record_count": len(records),
    }, records


def detect_header_row(rows: list[list[Any]], max_scan_rows: int = 10) -> int:
    aliases = {alias for names in FIELD_ALIASES.values() for alias in names}
    best_index = 0
    best_score = -1
    for index, row in enumerate(rows[:max_scan_rows]):
        score = sum(1 for value in row if normalize_text(value) in aliases)
        if score > best_score:
            best_index = index
            best_score = score
    return best_index


def map_headers(headers: list[str | None]) -> dict[str, int]:
    normalized = {header: idx for idx, header in enumerate(headers) if header}
    mapping: dict[str, int] = {}
    for standard_field, aliases in FIELD_ALIASES.items():
        for alias in aliases:
            if alias in normalized:
                mapping[standard_field] = normalized[alias]
                break
    return mapping


def build_raw_row(headers: list[str | None], row: list[Any]) -> dict[str, Any]:
    raw: dict[str, Any] = {}
    for idx, value in enumerate(row):
        key = headers[idx] if idx < len(headers) and headers[idx] else f"unnamed_column_{idx + 1}"
        raw[key] = value
    return raw


def build_source_cells(
    source_name: str,
    headers: list[str | None],
    row: list[Any],
    field_map: dict[str, int],
    row_index: int,
) -> dict[str, SourceCell]:
    sources: dict[str, SourceCell] = {}
    for field, idx in field_map.items():
        if idx >= len(row):
            continue
        sources[field] = SourceCell(
            sheet=source_name,
            row=row_index,
            column=idx + 1,
            header=headers[idx] if idx < len(headers) else None,
            raw_value=row[idx],
        )
    return sources


def get_field_value(field_map: dict[str, int], field: str, row: list[Any]) -> Any:
    idx = field_map.get(field)
    if idx is None or idx >= len(row):
        return None
    return row[idx]


def empty_report(source_name: str) -> dict[str, Any]:
    return {
        "source": source_name,
        "row_count": 0,
        "header_row": None,
        "headers": [],
        "field_map": {},
        "record_count": 0,
    }


def build_payload(
    input_path: Path,
    archived_source: Path | None,
    records: list[VoucherRecord],
    mapping_reports: list[dict[str, Any]],
    include_debug: bool,
) -> dict[str, Any]:
    payload: dict[str, Any] = {
        "schema": "VoucherRecord",
        "schema_version": "2.0",
        "purpose": "research_project_financial_acceptance_analysis",
        "source_document_type": input_path.suffix.lower().lstrip("."),
        "record_id_rule": "sha1(source_file_name|source_section|row|voucher_no|voucher_date)[:16]",
        "record_count": len(records),
        "records": [record.to_agent_dict() for record in records],
    }
    if include_debug:
        payload["debug"] = {
            "source_file": str(input_path),
            "archived_source_file": str(archived_source) if archived_source else None,
            "generated_at": datetime.now().isoformat(timespec="seconds"),
            "mapping_report": mapping_reports,
            "records": [record.to_debug_dict() for record in records],
        }
    return payload


def write_json(payload: dict[str, Any], output_path: Path) -> None:
    output_path.parent.mkdir(parents=True, exist_ok=True)
    output_path.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")


def build_arg_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(description="Convert Excel, Word, or PDF finance voucher files to VoucherRecord JSON.")
    parser.add_argument("input", type=Path, help="Path to the source .xlsx, .docx, or .pdf file.")
    parser.add_argument("-o", "--output", type=Path, default=Path("data/processed/voucher_records.json"))
    parser.add_argument("--archive-dir", type=Path, default=Path("data/raw"))
    parser.add_argument("--include-debug", action="store_true", help="Include source cells, raw rows, and mapping details.")
    return parser


def main() -> None:
    parser = build_arg_parser()
    args = parser.parse_args()
    payload = parse_financial_document(args.input, raw_archive_dir=args.archive_dir, include_debug=args.include_debug)
    write_json(payload, args.output)
    print(f"Wrote {payload['record_count']} VoucherRecord rows to {args.output}")


if __name__ == "__main__":
    main()
